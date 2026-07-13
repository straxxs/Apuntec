#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generador del Documento Informativo Kiroku.docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for lv in range(1, 4):
    doc.styles[f"Heading {lv}"].font.color.rgb = RGBColor(0x26, 0x32, 0x38)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="263238"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t


# ═══ PORTADA ═══
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KIROKU")
r.font.size = Pt(28)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documento Informativo y de Presentación de Proyecto")
r.font.size = Pt(14)

doc.add_paragraph("")

add_table(doc,
    ["Campo", "Valor"],
    [
        ["Patrocinador (Sponsor)", "Equipo directivo - Escuela Fragata Libertad N°21"],
        ["Gerente de Proyecto", "León Veraldi Rita"],
        ["Fecha", "13/07/2026"],
        ["Versión", "1.0"],
    ],
)

doc.add_page_break()


# ═══ 1. RESUMEN EJECUTIVO ═══
doc.add_heading("1. Resumen Ejecutivo", level=1)
doc.add_paragraph(
    "Kiroku es una plataforma web colaborativa diseñada para estudiantes de la Escuela "
    "Técnica Fragata Libertad N°21. Su objetivo es centralizar los apuntes, resúmenes y "
    "trabajos prácticos de todas las materias curriculares en un solo lugar, clasificados "
    "por día y materia. Cuando un estudiante falta, puede consultar directamente la "
    "plataforma para ver qué se hizo en clase, eliminando la dependencia de grupos "
    "informales de WhatsApp."
)
doc.add_paragraph(
    "El proyecto se desarrolló en 10 semanas bajo metodología Scrum, con un equipo de "
    "9 integrantes, utilizando Python/Flask en backend, JavaScript vanilla en frontend, "
    "MySQL/MariaDB como base de datos, y desplegado en Render (producción) con TiDB "
    "Cloud Serverless como base de datos en la nube."
)


# ═══ 2. CONTEXTO Y JUSTIFICACIÓN ═══
doc.add_heading("2. Contexto y Justificación", level=1)

doc.add_heading("Situación Actual", level=2)
doc.add_paragraph(
    "En muchas escuelas secundarias, cuando un estudiante falta por enfermedad, "
    "actividades deportivas, problemas personales o viajes, pierde el acceso a los "
    "contenidos trabajados en clase. La recuperación de estos apuntes depende de pedir "
    "fotos por grupos de WhatsApp o contactar a compañeros de manera individual, lo que "
    "genera:"
)
for item in [
    "Pérdida o fuga de información.",
    "Apuntes incompletos o de mala calidad visual.",
    "Desorganización general del material.",
    "Desigualdad en el acceso al contenido académico.",
    "Menor rendimiento académico de los estudiantes afectados.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Necesidad del Proyecto", level=2)
doc.add_paragraph(
    "Kiroku resuelve la necesidad concreta de los estudiantes de contar con un repositorio "
    "único, organizado y accesible de apuntes escolares. La plataforma elimina la "
    "fragmentación de información que generan los grupos de WhatsApp y garantiza que "
    "ningún estudiante se quede atrás por haber faltado a clase."
)

doc.add_heading("Alineación Estratégica", level=2)
doc.add_paragraph(
    "El proyecto se alinea con los objetivos de transformación digital e innovación "
    "educativa de la Escuela Técnica Fragata Libertad N°21, promoviendo el uso de "
    "tecnología para mejorar la experiencia de aprendizaje y la colaboración entre "
    "estudiantes."
)


# ═══ 3. OBJETIVOS ═══
doc.add_heading("3. Objetivos del Proyecto", level=1)

doc.add_heading("3.1 Objetivo General", level=2)
doc.add_paragraph(
    "Diseñar e implementar una plataforma web que facilite el acceso democrático y "
    "organizado a los materiales de estudio compartidos por los propios estudiantes."
)

doc.add_heading("3.2 Objetivos Específicos", level=2)
for o in [
    "Centralizar los apuntes de todas las materias curriculares en un solo lugar.",
    "Facilitar la recuperación rápida de contenidos ante inasistencias.",
    "Fomentar el compañerismo y la colaboración activa entre alumnos.",
    "Reducir la dependencia de los grupos informales de mensajería (WhatsApp/Telegram).",
    "Mejorar la organización académica del alumnado.",
]:
    doc.add_paragraph(o, style="List Bullet")

doc.add_heading("3.3 Indicadores de Éxito (KPIs)", level=2)
add_table(doc,
    ["Indicador", "Línea Base", "Meta", "Plazo"],
    [
        ["Cantidad de usuarios registrados", "0", "≥ 100 estudiantes activos", "Semana 10"],
        ["Apuntes subidos a la plataforma", "0", "≥ 50 apuntes publicados", "Semana 10"],
        ["Tasa de adopción del curso", "0%", "≥ 60% del curso registrado", "Semana 8"],
        ["Tiempo de carga de la plataforma", "N/A", "< 3 segundos", "Semana 5"],
        ["Disponibilidad del sistema", "N/A", "≥ 99% en horario escolar", "Semana 10"],
    ],
)


# ═══ 4. ALCANCE ═══
doc.add_heading("4. Alcance del Proyecto", level=1)

doc.add_heading("4.1 Dentro del Alcance (In Scope)", level=2)
for item in [
    "Sistema de registro, inicio de sesión y recuperación de contraseña.",
    "Subida, descarga y visualización de apuntes (PDF, imágenes, videos) con drag & drop.",
    "Clasificación de contenido por materia, curso, fecha y autor.",
    "Búsqueda avanzada con filtros (materia, curso, fecha, palabras clave).",
    "Sistema de valoración con estrellas (1-5) y \"me gusta\".",
    "Sistema de guardados para apuntes favoritos.",
    "Aprobación de contenido por moderadores antes de publicación general.",
    "Panel de administración con gestión de usuarios (bloquear, roles, eliminar).",
    "Dashboard de estadísticas con gráficos (Chart.js).",
    "Roles de usuario: Alumno, Moderador y Administrador.",
    "Código de invitación para unirse a cursos.",
    "Sistema de auditoría con logs de actividad.",
    "Diseño responsive optimizado para móvil.",
    "Efectos de sonido sutiles (Web Audio API).",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.2 Fuera del Alcance (Out of Scope)", level=2)
for item in [
    "Aplicación móvil nativa (Android/iOS) – prevista para V2.0.",
    "Chat interno en tiempo real entre estudiantes.",
    "Motor de recomendación automática basado en IA.",
    "Sincronización con calendarios escolares.",
    "Generación automática de cuestionarios (quizzes).",
    "Video-llamadas o videoconferencias integradas.",
    "Servicio de email real (SMTP) para recuperación de contraseñas (simulado).",
    "Almacenamiento de archivos en la nube (AWS S3).",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.3 Principales Entregables", level=2)
add_table(doc,
    ["Entregable", "Descripción"],
    [
        ["Sprint 1", "Sistema de autenticación funcional (registro, login, JWT, bcrypt)."],
        ["Sprint 2", "Módulo de materias y cursos operativo (CRUD, códigos de invitación)."],
        ["Sprint 3", "Sistema central de apuntes (subida, descarga, aprobación, vista previa)."],
        ["Sprint 4", "Motor de búsqueda, valoraciones, guardados y efectos de sonido."],
        ["Sprint 5", "Panel de administración, estadísticas, auditoría y deploy en Render."],
    ],
)


# ═══ 5. INTERESADOS CLAVE ═══
doc.add_heading("5. Interesados Clave (Stakeholders)", level=1)
add_table(doc,
    ["Nombre / Área", "Rol", "Interés / Expectativa", "Influencia", "Canal"],
    [
        ["Profesor/a de la materia", "Cliente / Product Owner", "Aprobar entregas, definir requisitos", "Alta", "Presencial / WhatsApp"],
        ["Equipo de desarrollo (9)", "Ejecutores", "Desarrollar el producto", "Alta", "WhatsApp / Trello / GitHub"],
        ["Estudiantes de la escuela", "Usuarios finales", "Subir y consultar apuntes", "Alta", "Plataforma web"],
        ["Preceptores y autoridades", "Supervisores", "Supervisar uso adecuado", "Media", "Presencial"],
        ["Equipo directivo", "Sponsor", "Autorizar despliegue en la escuela", "Media", "Presencial"],
    ],
)


# ═══ 6. CRONOGRAMA ═══
doc.add_heading("6. Cronograma y Fases", level=1)
add_table(doc,
    ["Fase / Hito", "Descripción", "Fecha Inicio", "Fecha Fin"],
    [
        ["Inicio", "Definición del proyecto, equipo y herramientas", "15/06/2026", "16/06/2026"],
        ["Sprint 1", "Autenticación y base (registro, login, DB)", "17/06/2026", "24/06/2026"],
        ["Sprint 2", "Gestión académica (materias, cursos, roles)", "25/06/2026", "02/07/2026"],
        ["Sprint 3", "Core del negocio (subida, descarga, aprobación)", "03/07/2026", "10/07/2026"],
        ["Sprint 4", "Interacción (búsqueda, valoraciones, sonidos)", "11/07/2026", "15/07/2026"],
        ["Sprint 5", "Administración, estadísticas y deploy", "16/07/2026", "22/07/2026"],
    ],
)


# ═══ 7. EQUIPO Y RECURSOS ═══
doc.add_heading("7. Equipo y Recursos", level=1)

doc.add_heading("7.1 Equipo del Proyecto", level=2)
add_table(doc,
    ["Nombre", "Rol", "Dedicación"],
    [
        ["León Veraldi Rita", "Líder de proyecto / Scrum Master", "Full-time"],
        ["Matías Hayes", "Sublíder / Desarrollador Backend", "Full-time"],
        ["Vito Martin", "Desarrollador Frontend", "Full-time"],
        ["Thiago Montenegro", "Desarrollador Frontend", "Full-time"],
        ["Ramiro Tatone", "Desarrollador Backend", "Full-time"],
        ["Santino Moauro", "Administrador de Base de Datos (DBA)", "Full-time"],
        ["Federico Rojas", "Tester (QA)", "Full-time"],
        ["Juan Pablo Santisi", "Desarrollador Full Stack", "Partial-time"],
        ["Lucas Hamlin", "Desarrollador Full Stack", "Partial-time"],
    ],
)

doc.add_heading("7.2 Stack Tecnológico", level=2)
add_table(doc,
    ["Componente", "Tecnología"],
    [
        ["Backend", "Python 3.14.5 + Flask"],
        ["Base de datos", "MySQL / MariaDB 10.4.32 (local) / TiDB Cloud (producción)"],
        ["Autenticación", "bcrypt + JWT (HTTP-only)"],
        ["Frontend", "HTML5, CSS3, JavaScript vanilla"],
        ["Gráficos", "Chart.js v4"],
        ["Íconos", "Lucide Icons (SVG)"],
        ["Hosting", "Render (producción) + TiDB Cloud Serverless"],
        ["Control de versiones", "Git / GitHub"],
    ],
)

doc.add_heading("7.3 Presupuesto Estimado", level=2)
doc.add_paragraph(
    "El proyecto se desarrolla íntegramente con herramientas gratuitas y de código "
    "abierto, por lo que no requiere inversión económica directa:"
)
for r in [
    "Git/GitHub: gratuito (plan educativo).",
    "Trello: gratuito.",
    "Figma: gratuito (plan educativo).",
    "Hosting: Render (plan gratuito起步) + TiDB Cloud Serverless (gratuito起步).",
    "Dominio: no se requiere (URL de Render: *.onrender.com).",
    "Tiempo del equipo: 10 semanas, 9 integrantes.",
]:
    doc.add_paragraph(r, style="List Bullet")


# ═══ 8. RIESGOS ═══
doc.add_heading("8. Riesgos y Supuestos Principales", level=1)
doc.add_paragraph(
    "Para el detalle completo de la gestión de riesgos, ver la Matriz de Riesgos del "
    "proyecto (Matriz_de_Riesgo_Kiroku.docx). Principales riesgos identificados:"
)
add_table(doc,
    ["Riesgo / Supuesto", "Impacto Potencial", "Estrategia"],
    [
        ["Baja participación de integrantes", "Retraso en entregas de sprint", "Mitigar: Trello + seguimiento semanal"],
        ["Problemas de conectividad", "Impide acceso a herramientas", "Aceptar: documentar avances offline"],
        ["Curva de aprendizaje tecnologías nuevas", "Retraso en desarrollo inicial", "Mitigar: capacitación en Sprint 1"],
        ["Desacuerdos en el equipo", "Bloqueo de decisiones", "Evitar: consenso + escalación al PO"],
        ["Cambios de requisitos", "Scope creep", "Evitar: contrato firmado + control de cambios"],
        ["Pérdida de datos", "Pérdida de apuntes subidos", "Mitigar: backups de DB + git"],
    ],
)


# ═══ 9. COMUNICACIÓN ═══
doc.add_heading("9. Comunicación y Gobernanza", level=1)

doc.add_heading("9.1 Estructura de Gobernanza", level=2)
doc.add_paragraph(
    "El comité de dirección está conformado por el profesor/a de la materia (Product Owner) "
    "y el líder del equipo (Scrum Master). Las reuniones de revisión de Sprint se realizan "
    "cada 2 semanas. Las decisiones se toman por consenso del equipo, con escalación al "
    "Product Owner en caso de desacuerdos no resueltos."
)

doc.add_heading("9.2 Plan de Comunicación", level=2)
add_table(doc,
    ["Reporte / Reunión", "Audiencia", "Frecuencia", "Responsable"],
    [
        ["Reunión diaria (Daily)", "Equipo completo", "Diaria (10 min)", "Líder / Scrum Master"],
        ["Revisión de Sprint", "Equipo + PO", "Cada 2 semanas", "Scrum Master"],
        ["Reporte de estado", "Profesor / Cliente", "Semanal", "Líder"],
        ["Comunicación por WhatsApp", "Equipo completo", "Diaria (asíncrona)", "Todos"],
        ["Tablero Trello", "Equipo completo", "Actualización continua", "Cada integrante"],
    ],
)


# ═══ GUARDAR ═══
output_dir = os.path.join(os.path.dirname(__file__), "Documentación")
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "Documento_Informativo_Kiroku.docx")
doc.save(output_path)
print(f"Documento generado: {output_path}")
