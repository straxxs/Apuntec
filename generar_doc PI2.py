#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador del documento Documentacion PI2.docx
Reconstruye el documento completo desde cero con todas las secciones.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Estilos base ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x26, 0x32, 0x38)

# ── Helper: tabla con header oscuro ──
def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="263238"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t


# ════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════
doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Escuela Técnica Fragata Libertad N°21"')
r.font.size = Pt(14)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DOCUMENTACIÓN")
r.font.size = Pt(22)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROYECTO INFORMÁTICO 2")
r.font.size = Pt(16)
r.bold = True

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma Colaborativa de Apuntes para Estudiantes")
r.font.size = Pt(13)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Nombre del Proyecto: KIROKU")
r.font.size = Pt(12)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Ningún estudiante se queda atrás."')
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Equipo: León Veraldi Rita (líder), Vito Martin, Matías Hayes (sublíder), "
    "Thiago Montenegro, Ramiro Tatone, Santino Moauro, Federico Rojas, "
    "Juan Pablo Santisi, y Lucas Hamlin"
)
r.font.size = Pt(10)

doc.add_page_break()


# ════════════════════════════════════════════
# LÍMITES Y NORMAS DEL EQUIPO
# ════════════════════════════════════════════
doc.add_heading("Límites y Normas del Equipo de Trabajo", level=1)

normas = [
    "Todos los integrantes deberán participar activamente en el desarrollo del proyecto y cumplir con las tareas asignadas.",
    "Cada miembro será responsable de mantener actualizada su tarea en Trello, indicando correctamente su estado de avance.",
    "Las tareas deberán completarse dentro de los plazos establecidos para evitar retrasos en el proyecto.",
    "Toda tarea deberá encontrarse en una de las siguientes listas: Pendiente, En proceso, En revisión o Finalizado.",
    "Ninguna tarea podrá marcarse como finalizada sin haber sido completada y revisada correctamente.",
    "Los integrantes deberán comunicarse de manera respetuosa y colaborar entre sí cuando sea necesario.",
    "Cualquier cambio importante en la planificación o en las tareas deberá ser informado al líder y al resto del equipo.",
    "Los integrantes que finalicen sus tareas podrán brindar apoyo a otros compañeros que lo necesiten.",
    "Se deberán conservar las evidencias del trabajo realizado, incluyendo capturas de pantalla y registros de avances.",
    "El líder del equipo, León Veraldi Rita, será el encargado de coordinar las actividades, supervisar el progreso y resolver inconvenientes organizativos.",
    "Ante desacuerdos o conflictos, se buscará llegar a una solución mediante el diálogo y el consenso del equipo.",
    "El objetivo principal será completar el proyecto de manera organizada, colaborativa y dentro de los tiempos establecidos.",
]
for n in normas:
    doc.add_paragraph(n, style="List Bullet")


# ════════════════════════════════════════════
# BITÁCORA DE AVANCES
# ════════════════════════════════════════════
doc.add_heading("Bitácora de Avances", level=1)

doc.add_heading("22/06/2026", level=3)
doc.add_paragraph(
    "Empezamos a desarrollar el tablero del Trello para la actividad de Producción. "
    "El día de hoy asistimos todos menos Lucas Hamlin y Juan Santisi. En equipo fuimos "
    "añadiendo las tarjetas, las tareas, los miembros, los responsables, el nivel de "
    "prioridad y las fechas límites."
)
doc.add_paragraph(
    "También el día de hoy elegí a Matías Hayes como el sublíder del equipo."
)
doc.add_paragraph(
    "Durante la realización del trabajo se presentaron algunas dificultades: "
    "coordinación entre los integrantes del equipo, distribución inicial de "
    "responsabilidades, organización de los tiempos de trabajo, y seguimiento de las "
    "tareas pendientes para cumplir con las fechas establecidas. Estas dificultades "
    "pudieron resolverse gracias al uso de Trello y a la comunicación constante entre "
    "los integrantes."
)

doc.add_heading("24/06/2026", level=3)
doc.add_paragraph(
    "El día de hoy definimos con el equipo entero el proyecto a llevar a cabo para "
    "entregar antes de las vacaciones de invierno. El proyecto se trata de una página "
    "en la cual los alumnos que asistan a clases puedan subir lo que hicieron ese día, "
    "así los que faltaron puedan consultar directamente lo que hicieron sin tener que "
    "pedirlo (ya que muchas veces se olvidan de pedirlo o si lo piden no lo pasan). "
    "Este proyecto facilitará a los alumnos acceder a lo hecho cada día del año ya que "
    "se clasificará por días y por materias."
)

doc.add_page_break()


# ════════════════════════════════════════════
# 1. PROBLEMÁTICA
# ════════════════════════════════════════════
doc.add_heading("1. Problemática", level=1)
doc.add_paragraph(
    "En muchas escuelas secundarias, cuando un estudiante falta por enfermedad, "
    "actividades deportivas, problemas personales o viajes, pierde el acceso a los "
    "contenidos trabajados en clase. La recuperación de estos apuntes depende de pedir "
    "fotos por grupos de WhatsApp o contactar a compañeros de manera individual, lo que "
    "genera las siguientes dificultades:"
)
for d in [
    "Pérdida o fuga de información.",
    "Apuntes incompletos, fragmentados o de mala calidad visual.",
    "Desorganización general del material.",
    "Desigualdad en el acceso al contenido académico.",
    "Menor rendimiento académico de los estudiantes afectados.",
]:
    doc.add_paragraph(d, style="List Bullet")


# ════════════════════════════════════════════
# 2. SOLUCIÓN PROPUESTA
# ════════════════════════════════════════════
doc.add_heading("2. Solución Propuesta", level=1)
doc.add_paragraph(
    "Desarrollar e implementar una plataforma web donde los estudiantes puedan subir, "
    "organizar y compartir apuntes, resúmenes, trabajos prácticos y material de estudio "
    "clasificado por materia. La plataforma funcionará como la principal biblioteca "
    "digital colaborativa de la escuela."
)
doc.add_paragraph("Nombre del Proyecto: KIROKU").runs[0].bold = True
doc.add_paragraph('"Ningún estudiante se queda atrás."').runs[0].italic = True


# 2.1 Objetivos
doc.add_heading("2.1. Objetivos", level=2)
doc.add_heading("Objetivo General", level=3)
doc.add_paragraph(
    "Diseñar e implementar una plataforma web que facilite el acceso democrático y "
    "organizado a los materiales de estudio compartidos por los propios estudiantes."
)
doc.add_heading("Objetivos Específicos", level=3)
for o in [
    "Centralizar los apuntes de todas las materias curriculares en un solo lugar.",
    "Facilitar la recuperación rápida de contenidos ante inasistencias.",
    "Fomentar el compañerismo y la colaboración activa entre alumnos.",
    "Reducir la dependencia de los grupos informales de mensajería (WhatsApp/Telegram).",
    "Mejorar la organización académica del alumnado.",
]:
    doc.add_paragraph(o, style="List Bullet")


# ════════════════════════════════════════════
# 3. PERFILES DE USUARIOS
# ════════════════════════════════════════════
doc.add_heading("3. Perfiles de Usuarios", level=1)
doc.add_paragraph(
    "El sistema cuenta con tres niveles de acceso claramente definidos:"
)

doc.add_heading("Alumno (Usuario Estándar)", level=3)
for a in [
    "Registrarse e iniciar sesión de forma segura.",
    "Subir, buscar, valorar y descargar apuntes.",
    "Calificar y dar \"me gusta\" al material de sus compañeros.",
    "Guardar apuntes favoritos.",
]:
    doc.add_paragraph(a, style="List Bullet")

doc.add_heading("Delegado o Moderador", level=3)
doc.add_paragraph("Posee los permisos de Alumno, más:")
for m in [
    "Aprobar contenido nuevo antes de su publicación general.",
    "Eliminar material incorrecto, repetido o inapropiado.",
    "Gestionar y auditar materias de su curso.",
    "Ver estadísticas del curso.",
    "Subir apuntes que se publican directamente (sin revisión).",
]:
    doc.add_paragraph(m, style="List Bullet")

doc.add_heading("Administrador", level=3)
for a in [
    "Gestionar usuarios (bloquear/activar, asignar roles).",
    "Ver panel de estadísticas generales de la plataforma.",
    "Gestionar la creación de nuevos cursos.",
    "Eliminar usuarios y cursos.",
]:
    doc.add_paragraph(a, style="List Bullet")


# ════════════════════════════════════════════
# 4. FUNCIONALIDADES DEL SISTEMA
# ════════════════════════════════════════════
doc.add_heading("4. Funcionalidades del Sistema (Módulos)", level=1)

modulos = [
    ("Módulo de Usuarios",
     "Registro/Login con nombre de usuario o email, recuperación de contraseña por email "
     "(simulado), gestión de perfil (avatar, nombre), selección de curso por código de "
     "invitación, y sistema de roles (alumno, moderador, administrador)."),
    ("Módulo de Materias",
     "CRUD de materias por curso. Los moderadores pueden crear, editar y eliminar materias "
     "dentro de su curso asignado."),
    ("Módulo de Apuntes",
     "Estructura de cada apunte: título, materia, curso, autor, fecha de subida, archivos "
     "(PDF, imágenes, videos) con drag & drop, y descripción breve. Los moderadores suben "
     "apuntes que se publican directamente; los alumnos suben apuntes que quedan pendientes "
     "de aprobación."),
    ("Módulo de Aprobación",
     "Los moderadores revisan apuntes pendientes y los aprueban o rechazan con observaciones. "
     "Solo el contenido aprobado es visible para todos los alumnos."),
    ("Módulo de Búsqueda",
     "Filtros de búsqueda avanzados por materia, curso, fecha, autor y palabras clave."),
    ("Módulo de Valoración",
     "Sistema de estrellas (1-5), botón \"Me Gusta\", sistema de guardados para apuntes "
     "favoritos, y página \"Mis guardados\" en el perfil."),
    ("Módulo de Estadísticas",
     "Dashboard con gráficos (Chart.js): métricas generales, ranking de colaboradores, "
     "apuntes más descargados. Estadísticas por curso: resumen, heatmap de actividad, "
     "ranking de colaboradores, top valorados."),
    ("Módulo de Administración",
     "Panel de administración con gestión de usuarios (tabla con nombre, rol, estado, curso, "
     "acciones de bloquear/activar, ascender/descender, eliminar) y gestión de cursos "
     "(tabla con año, división, creador, acción de eliminar)."),
    ("Módulo de Auditoría",
     "Sistema de logs con usuario, acción, timestamp. Se registran 22 eventos: login, "
     "registro, subida de apuntes, aprobación, eliminación, cambio de rol, bloqueo, etc."),
    ("Módulo de Seguridad",
     "Autenticación con bcrypt (hash) y JWT (tokens HTTP-only), headers de seguridad "
     "(CSP, X-Frame-Options, etc.), validación robusta de contraseñas (longitud, "
     "complejidad, fortaleza), protección XSS con escapeHtml centralizado, y cookies HttpOnly."),
]
for nombre, desc in modulos:
    doc.add_heading(nombre, level=3)
    doc.add_paragraph(desc)


# ════════════════════════════════════════════
# 5. ARQUITECTURA Y TECNOLOGÍAS
# ════════════════════════════════════════════
doc.add_heading("5. Arquitectura y Tecnologías", level=1)
doc.add_paragraph(
    "El sistema sigue una arquitectura Cliente-Servidor (Frontend/Backend) con las "
    "siguientes tecnologías:"
)

add_table(doc,
    ["Componente", "Tecnología"],
    [
        ["Backend", "Python 3.14.5 + Flask"],
        ["Base de datos", "MySQL / MariaDB 10.4.32 (XAMPP local) / TiDB Cloud (producción)"],
        ["Autenticación", "bcrypt (hash de contraseñas) + JWT (sesiones HTTP-only)"],
        ["Frontend", "HTML5, CSS3, JavaScript vanilla"],
        ["Gráficos", "Chart.js v4"],
        ["Audio", "Web Audio API (sonidos generados en código)"],
        ["Íconos", "Lucide Icons (SVG)"],
        ["Servidor", "Flask dev server (desarrollo) / Gunicorn (producción)"],
        ["Hosting", "Render (producción) + TiDB Cloud Serverless (DB)"],
        ["Control de versiones", "Git / GitHub"],
        ["Gestión de tareas", "Trello"],
        ["Diseño UI", "Figma"],
    ],
)
doc.add_paragraph("")

doc.add_heading("Estructura del Proyecto", level=2)
doc.add_paragraph(
    "El proyecto está organizado en la siguiente estructura de carpetas:"
)
estructura = """Proyecto Mitingay/
├── App/
│   ├── app.py                  # Rutas principales (Flask)
│   ├── wsgi.py                 # Entry point para Gunicorn (Render)
│   ├── db/
│   │   └── conexion.py         # Conexión a MySQL (soporta DATABASE_URL)
│   ├── modulos/
│   │   ├── auth.py             # Registro, login, JWT, bcrypt
│   │   ├── usuarios.py         # CRUD usuarios, roles, estado
│   │   ├── cursos.py           # CRUD cursos, códigos de invitación
│   │   ├── materias.py         # CRUD materias
│   │   ├── apuntes.py          # CRUD apuntes y archivos
│   │   ├── valoraciones.py     # Estrellas, guardados, me gusta
│   │   ├── busqueda.py         # Búsqueda avanzada
│   │   ├── estadisticas.py     # Métricas, rankings, stats por curso
│   │   ├── recuperacion.py     # Tokens de recuperación
│   │   ├── validacion.py       # Validación de contraseña, usuario, email
│   │   ├── auditoria.py        # Logs de auditoría
│   │   └── profesores.py       # Gestión de profesores
│   ├── templates/              # HTML (Jinja2)
│   ├── static/
│   │   ├── css/                # Estilos (styles.css, panel.css)
│   │   ├── js/                 # Módulos JS (17 archivos)
│   │   ├── img/logo.png        # Logo KIROKU
│   │   └── uploads/            # Avatares y archivos subidos
│   ├── mitin.sql               # Script de inicialización de DB
│   └── requirements.txt        # Dependencias Python
├── Documentación/              # Documentos del proyecto
└── README.md"""
p = doc.add_paragraph()
r = p.add_run(estructura)
r.font.size = Pt(9)
r.font.name = "Consolas"


doc.add_page_break()


# ════════════════════════════════════════════
# 6. METODOLOGÍA Y ROLES
# ════════════════════════════════════════════
doc.add_heading("6. Metodología de Trabajo y Roles del Equipo", level=1)
doc.add_paragraph(
    "El proyecto se desarrolló bajo la Metodología Ágil Scrum, con una duración total "
    "de 10 semanas, divididas en 5 Sprints (2 semanas por cada iteración)."
)

doc.add_heading("Roles Designados", level=2)
roles = [
    ("Product Owner", "Define requisitos y prioriza funcionalidades (Profesor tutor). Aprueba las entregas de valor."),
    ("Scrum Master / Líder", "Coordina ceremonias ágiles, resuelve bloqueos técnicos y vela por el cumplimiento de la metodología. Líder: León Veraldi Rita."),
    ("Sublíder", "Apoya al líder en la coordinación del equipo. Sublíder: Matías Hayes."),
    ("Desarrollador Frontend", "Responsable de diseño visual, maquetación, formularios y navegación de usuario."),
    ("Desarrollador Backend", "Responsable de APIs, lógica de negocio, autenticación y seguridad."),
    ("Administrador de Base de Datos (DBA)", "Diseño del modelo Entidad-Relación, consultas SQL y optimización de datos."),
    ("Tester (QA)", "Redacción de casos de prueba, reporte de errores y validación de calidad antes de las entregas."),
]
for nombre, desc in roles:
    p = doc.add_paragraph()
    r = p.add_run(f"{nombre}: ")
    r.bold = True
    p.add_run(desc)


# ════════════════════════════════════════════
# 7. PRODUCT BACKLOG
# ════════════════════════════════════════════
doc.add_heading("7. Product Backlog (Historias de Usuario)", level=1)

add_table(doc,
    ["ID", "Historia de Usuario"],
    [
        ["HU1", "Registrarme en la plataforma"],
        ["HU2", "Iniciar sesión (Login)"],
        ["HU3", "Subir apuntes y documentos"],
        ["HU4", "Buscar y filtrar apuntes"],
        ["HU5", "Descargar apuntes a mi dispositivo"],
        ["HU6", "Valorar y calificar apuntes de otros"],
        ["HU7", "Gestionar materias y categorizaciones"],
        ["HU8", "Administrar roles y usuarios"],
        ["HU9", "Ver métricas y estadísticas del sistema"],
        ["HU10", "Recuperar contraseña de acceso"],
        ["HU11", "Guardar apuntes en favoritos"],
        ["HU12", "Aprobar o rechazar contenido (moderador)"],
        ["HU13", "Unirse a un curso con código de invitación"],
        ["HU14", "Personalizar mi perfil (avatar, curso)"],
    ],
)


# ════════════════════════════════════════════
# 8. CRONOGRAMA DE SPRINTS
# ════════════════════════════════════════════
doc.add_heading("8. Cronograma de Sprints", level=1)

sprints = [
    ("Sprint 1: Autenticación y Base", "Semanas 1-2",
     "Sistema de autenticación funcional (registro, login, JWT, bcrypt), modelo de base de datos, "
     " wireframes y pantallas de login/registro."),
    ("Sprint 2: Gestión Académica", "Semanas 3-4",
     "Gestión de materias y cursos (CRUD), navegación dinámica, sistema de roles y permisos, "
     "selección de curso por código de invitación."),
    ("Sprint 3: Core del Negocio", "Semanas 5-6",
     "Sistema central de apuntes: subida de archivos (PDF, imágenes, videos) con drag & drop, "
     "aprobación de contenido por moderadores, descarga directa, vista previa de archivos."),
    ("Sprint 4: Interacción y Búsqueda", "Semanas 7-8",
     "Motor de búsqueda avanzada, sistema de valoraciones (estrellas, me gusta), guardados de "
     "favoritos, efectos de sonido, doodles decorativos."),
    ("Sprint 5: Administración y Cierre", "Semanas 9-10",
     "Panel de administración (gestión de usuarios y cursos), dashboard de estadísticas con "
     "Chart.js, sistema de auditoría, modo responsive para móvil, deploy en Render."),
]
for nombre, periodo, desc in sprints:
    doc.add_heading(f"{nombre} ({periodo})", level=3)
    doc.add_paragraph(desc)


doc.add_page_break()


# ════════════════════════════════════════════
# 9. REQUERIMIENTOS FUNCIONALES Y NO FUNCIONALES
# ════════════════════════════════════════════
doc.add_heading("9. Requerimientos Funcionales y No Funcionales", level=1)
doc.add_paragraph(
    "A continuación se detallan los requerimientos del contrato de proyecto y su estado "
    "de implementación:"
)

doc.add_heading("Requerimientos Funcionales", level=2)
add_table(doc,
    ["ID", "Requerimiento", "Estado", "Detalle"],
    [
        ["RF-01", "Registro de usuario", "Cumplido",
         "Registro con nombre, email y contraseña. Validación de fortaleza. Un solo curso por usuario."],
        ["RF-02", "Inicio de sesión", "Cumplido",
         "Login con nombre de usuario O email. Rechaza usuarios bloqueados."],
        ["RF-03", "Recuperación de contraseña", "Simulado",
         "Página de recuperación con tokens y formulario de reseteo. No envía emails reales (sin SMTP)."],
        ["RF-04", "Subir apuntes", "Cumplido",
         "Upload de PDF, imágenes, videos. Drag & drop. Requiere archivo. Moderadores publican directo."],
        ["RF-05", "Buscar y filtrar", "Cumplido",
         "Búsqueda avanzada por materia, fecha, autor y palabras clave."],
        ["RF-06", "Descargar apuntes", "Cumplido",
         "Descarga directa de archivos con ruta correcta."],
        ["RF-07", "Valorar apuntes", "Cumplido",
         "Sistema de estrellas (1-5), botón Me Gusta, guardados en favoritos."],
        ["RF-08", "Aprobar contenido", "Cumplido",
         "Moderadores aprueban/rechazan apuntes pendientes con observaciones."],
        ["RF-09", "Gestionar usuarios", "Cumplido",
         "Admin puede bloquear/activar, ascender/descender roles, eliminar usuarios."],
        ["RF-10", "Ver estadísticas", "Cumplido",
         "Dashboard con Chart.js: métricas, rankings, heatmap de actividad por curso."],
    ],
)
doc.add_paragraph("")

doc.add_heading("Requerimientos No Funcionales", level=2)
add_table(doc,
    ["ID", "Requerimiento", "Estado", "Detalle"],
    [
        ["RNF-01", "Respuesta < 3 segundos", "Cumplido",
         "Flask + MySQL local responden en < 1s. En Render (producción) depende del plan gratuito."],
        ["RNF-02", "bcrypt + JWT", "Cumplido",
         "Contraseñas hasheadas con bcrypt. Sesiones manejadas con JWT HTTP-only."],
        ["RNF-03", "Disponibilidad 99%", "Depende de hosting",
         "Depende del tier de Render. En desarrollo local es 100%."],
        ["RNF-05", "Escalabilidad", "Depende de hosting",
         "Arquitectura stateless (JWT) permite escalar horizontalmente."],
        ["RNF-08", "Auditoría / trazabilidad", "Cumplido",
         "Sistema completo de logs: 22 eventos, tabla audit_log, funciones registrar_accion() y listar_auditoria()."],
    ],
)


doc.add_page_break()


# ════════════════════════════════════════════
# 10. PRESUPUESTO ESTIMADO
# ════════════════════════════════════════════
doc.add_heading("10. Presupuesto Estimado", level=1)
doc.add_paragraph(
    "El proyecto se desarrolla íntegramente con herramientas gratuitas y de código "
    "abierto, por lo que no requiere inversión económica directa. Los recursos "
    "estimados son:"
)
for r in [
    "Git/GitHub: gratuito (plan educativo).",
    "Trello: gratuito.",
    "Figma: gratuito (plan educativo).",
    "Hosting: Render (plan gratuito para起步) + TiDB Cloud Serverless (gratuito para起步).",
    "Dominio: no se requiere (se usa la URL de Render: *.onrender.com).",
    "Tiempo del equipo: 10 semanas, 9 integrantes (recursos humanos del curso).",
]:
    doc.add_paragraph(r, style="List Bullet")


# ════════════════════════════════════════════
# 11. INTERESADOS CLAVE
# ════════════════════════════════════════════
doc.add_heading("11. Interesados Clave (Stakeholders)", level=1)

add_table(doc,
    ["Nombre / Área", "Rol", "Interés / Expectativa", "Nivel de Influencia"],
    [
        ["Profesor/a de la materia", "Cliente / Product Owner", "Aprobar entregas, definir requisitos", "Alta"],
        ["Equipo de desarrollo (9 integrantes)", "Ejecutores", "Desarrollar el producto", "Alta"],
        ["Estudiantes de la escuela", "Usuarios finales", "Usar la plataforma para subir/consultar apuntes", "Alta"],
        ["Preceptores y autoridades", "Supervisores", "Supervisar el uso adecuado de la plataforma", "Media"],
        ["Equipo directivo", "Sponsor", "Autorizar el despliegue en la escuela", "Media"],
    ],
)


# ════════════════════════════════════════════
# 12. MATRIZ DE RIESGOS
# ════════════════════════════════════════════
doc.add_heading("12. Matriz de Riesgos", level=1)
doc.add_paragraph(
    "Se identificaron 8 riesgos principales para el proyecto, clasificados por "
    "probabilidad (1-5) e impacto (1-5). La severidad (P x I) determina el nivel de "
    "riesgo: Bajo (1-4), Medio (5-9), Alto (10-15) o Crítico (16-25). Cada riesgo "
    "cuenta con una estrategia de respuesta (Evitar, Mitigar, Transferir o Aceptar) y "
    "un responsable asignado."
)
doc.add_paragraph(
    "Ver el detalle completo en el documento separado: Matriz_de_Riesgo_Kiroku.docx"
).runs[0].italic = True


# ════════════════════════════════════════════
# 13. PLAN DE COMUNICACIÓN
# ════════════════════════════════════════════
doc.add_heading("13. Plan de Comunicación", level=1)

add_table(doc,
    ["Reporte / Reunión", "Audiencia", "Frecuencia", "Responsable"],
    [
        ["Reunión diaria (Daily)", "Equipo completo", "Diaria (10 min)", "Líder / Scrum Master"],
        ["Revisión de Sprint", "Equipo + Product Owner", "Cada 2 semanas", "Scrum Master"],
        ["Reporte de estado", "Profesor / Cliente", "Semanal", "Líder"],
        ["Comunicación por WhatsApp", "Equipo completo", "Diaria (asíncrona)", "Todos"],
        ["Tablero Trello", "Equipo completo", "Actualización continua", "Cada integrante"],
    ],
)


# ════════════════════════════════════════════
# 14. INDICADORES DE ÉXITO
# ════════════════════════════════════════════
doc.add_heading("14. Indicadores de Éxito (KPIs)", level=1)

add_table(doc,
    ["Indicador", "Línea Base", "Meta", "Plazo"],
    [
        ["Cantidad de usuarios registrados", "0", "≥ 100 estudiantes activos", "Semana 10"],
        ["Apuntes subidos a la plataforma", "0", "≥ 50 apuntes publicados", "Semana 10"],
        ["Tasa de adopción del curso", "0%", "≥ 60% del curso registrado", "Semana 8"],
        ["Tiempo de carga de la plataforma", "N/A", "< 3 segundos", "Semana 5"],
        ["Disponibilidad del sistema", "N/A", "≥ 99% en horario escolar", "Semana 10"],
    ],
)


doc.add_page_break()


# ════════════════════════════════════════════
# 15. PROBLEMÁTICAS DEL GRUPO Y SOLUCIONES
# ════════════════════════════════════════════
doc.add_heading("15. Problemáticas del Grupo y Soluciones", level=1)
doc.add_paragraph(
    "Durante el desarrollo del proyecto KIROKU, el equipo enfrentó diversas dificultades "
    "técnicas y organizativas. A continuación se detallan las principales problemáticas "
    "y cómo fueron resueltas:"
)

# -- Organizativas --
doc.add_heading("15.1. Problemáticas Organizativas", level=2)

problemas_org = [
    ("Coordinación de 9 integrantes",
     "Con un equipo de 9 personas, era difícil mantener a todos informados y alineados.",
     "Se implementó Trello como tablero de tareas con listas Pendiente/En proceso/En revisión/Finalizado. "
     "Se creó un grupo de WhatsApp para comunicación asíncrona diaria. Cada miembro era responsable "
     "de actualizar su tarjeta en Trello."),
    ("Distribución inicial de responsabilidades",
     "Al inicio no estaba claro quién hacía qué, lo que generaba duplicación de trabajo o tareas sin assigned.",
     "Se designó un sublíder (Matías Hayes) para apoyar al líder en la coordinación. Se asignaron "
     "roles específicos (Frontend, Backend, DBA, QA) a cada integrante según sus fortalezas."),
    ("Asistencia desigual a las reuniones",
     "No todos los integrantes asistían a todas las reuniones presenciales.",
     "Se documentaban las decisiones tomadas y se compartían por WhatsApp. Las tareas se asignaban "
     "de manera asíncrona en Trello para que cada quien pudiera avanzar a su ritmo."),
    ("Seguimiento de plazos",
     "Algunos integrantes no cumplían con las fechas límite de sus tareas.",
     "Se establecieron revisiones de sprint cada 2 semanas. El líder y sublíder hacían seguimiento "
     "semanal del avance y contactaban directamente a los miembros atrasados."),
]

for titulo, problema, solucion in problemas_org:
    doc.add_heading(titulo, level=3)
    p = doc.add_paragraph()
    r = p.add_run("Problema: ")
    r.bold = True
    p.add_run(problema)
    p = doc.add_paragraph()
    r = p.add_run("Solución: ")
    r.bold = True
    p.add_run(solucion)

# -- Técnicas --
doc.add_heading("15.2. Problemáticas Técnicas", level=2)

problemas_tec = [
    ("Cambio de nombre del proyecto",
     "El proyecto comenzó llamándose \"Appunti\", luego se renombró a \"Apuntec\" y finalmente "
     "a \"KIROKU\". Cada cambio requería actualizar templates, CSS, base de datos y documentación.",
     "Se hizo una búsqueda exhaustiva de todas las referencias al nombre anterior en el código "
     "fuente (HTML, CSS, JS, SQL) y se reemplazaron sistemáticamente. Se verificó con grep "
     "que no quedara ninguna referencia residual."),
    ("Stack tecnológico cambiante",
     "El contrato original planteaba \"Node.js o PHP\" y \"Bootstrap\". Se terminó usando "
     "Python + Flask + JavaScript vanilla por decisiones del equipo de desarrollo.",
     "Se actualizó la documentación del proyecto para reflejar el stack real utilizado. "
     "La arquitectura Flask + vanilla JS resultó ser ligera, rápida y adecuada para el alcance del proyecto."),
    ("Contraseña hasheada no matcheaba",
     "El usuario testuser2 no podía iniciar sesión porque el hash almacenado en la base de datos "
     "no coincidía con la contraseña \"1234\". Al verificar el hash con bcrypt, no matcheaba.",
     "Se reseteó manualmente el hash de la contraseña en la base de datos usando bcrypt.hashpw() "
     "desde Python para generar un hash válido de \"1234\" y actualizarlo con UPDATE."),
    ("Usuarios bloqueados podían iniciar sesión",
     "El botón de bloquear/desbloquear en el panel de administración funcionaba (cambiaba el "
     "estado en la BD), pero los usuarios bloqueados podían seguir iniciando sesión.",
     "Se identificó que la función login() en auth.py no verificaba el campo \"estado\". Se agregó "
     "la verificación: si estado == \"bloqueado\", se retorna un motivo \"bloqueado\" en vez de "
     "permitir el login. Se actualizó la ruta de login en app.py para mostrar \"Tu cuenta está bloqueada\". "
     "Además se agregó credentials: \"same-origin\" al fetch del admin.js y cache-busting al admin.html."),
    ("Merge conflict con branch remote",
     "Al intentar sincronizar el repositorio local con los cambios remotos del branch \"correcciones\", "
     "se produjo un merge conflict con múltiples archivos modificados.",
     "Se resolvió el merge conflict de manera manual, revisando cada conflicto archivo por archivo, "
     "manteniendo los cambios más recientes y completos. Se verificó que la aplicación funcionara "
     "correctamente después del merge."),
    ("XSS (Cross-Site Scripting) en múltiples puntos",
     "Se encontraron vulnerabilidades de XSS donde el contenido ingresado por los usuarios "
     "se renderizaba sin sanitizar, permitiendo inyección de scripts.",
     "Se implementó una función escapeHtml() centralizada en avatar.js que se reutiliza en "
     "admin.js, curso.js y otros módulos. Se sanitizan todos los datos que provienen de "
     "usuarios antes de renderizarlos en el DOM."),
    ("JWT stale después de cambiar avatar",
     "Cuando un usuario cambiaba su avatar, el token JWT seguía teniendo el avatar anterior. "
     "Al recargar la página, el avatar volvía al estado anterior hasta que se hacía logout/login.",
     "Se implementó un refresh automático del JWT después de cada cambio de perfil. La función "
     "de perfil ahora retorna un token nuevo con los datos actualizados, y el frontend lo "
     "almacena inmediatamente en localStorage."),
    ("DELETE cascade fallaba por FK constraints",
     "Al intentar eliminar un curso o usuario, la operación fallaba por restricciones de "
     "clave foránea (Foreign Key) en las tablas dependientes.",
     "Se revisó el modelo de base de datos y se agregaron cláusulas ON DELETE CASCADE a las "
     "foreign keys correspondientes. También se agregó eliminación en cadena manual en el "
     "módulo usuarios.py: al eliminar un usuario, se eliminan sus apuntes, valoraciones y "
     "guardados previamente."),
    ("Nombre de columna incompatible con MySQL",
     "La columna \"contraseña\" (con ñ) causaba problemas de encoding en MySQL. El motor de "
     "base de datos no procesaba correctamente los caracteres especiales en nombres de columnas.",
     "Se renombró la columna a \"contrasena\" (sin ñ) en el script SQL y en todos los módulos "
     "Python que hacían referencia a ella. Se verificó que todas las consultas SQL usaran el "
     "nuevo nombre."),
    ("Migración de emojis a íconos SVG",
     "Los emojis usados para íconos (🗂️, ⭐, 👍, etc.) se renderizaban de manera inconsistente "
     "entre navegadores y plataformas (Windows, Mac, Linux, móvil).",
     "Se creó un módulo centralizado icons.js con un objeto IC que contiene todas las funciones "
     "de renderizado de íconos usando Lucide Icons (SVG). Se reemplazaron TODOS los emojis del "
     "código fuente por llamadas a este módulo, asegurando renderizado consistente."),
    ("Dark mode incompleto",
     "Se había implementado un modo oscuro con toggle, pero quedó incompleto: algunos componentes "
     "no se veían bien, las variables CSS no cubrían todas las páginas, y el toggle no persistía.",
     "Se decidió remover el dark mode por completo para mantener consistencia visual. Se eliminaron "
     "las variables CSS del modo oscuro, el toggle del menú, y se consolidó el diseño en modo "
     "claro con la paleta post-it (celeste, amarillo, verde, rojo, violeta)."),
    ("Botones con rotación indeseada",
     "Los botones del sistema tenían una rotación CSS (transform: rotate) que se veía bien en "
     "desktop pero causaba problemas de alineación en móvil.",
     "Se eliminó la rotación de todos los botones (.btn-color, .btn-enviar) y se mantuvo solo "
     "el efecto press (translate + sombra) que funciona correctamente en todas las resoluciones."),
    ("Documento de documentación PI2 con duplicaciones",
     "El archivo Documentacion PI2.docx presentaba una duplicación masiva: las secciones 9.1 a 9.5 "
     "se repetían más de 25 veces, con tablas de Stakeholders, Plan de Comunicación e KPIs "
     "duplicadas 18 veces cada una. El documento tenía más de 50 tablas donde solo había 3 únicas.",
     "Se reconstruyó el documento completo desde cero usando python-docx, manteniendo el contenido "
     "válido, actualizando los datos obsoletos (nombre, stack tecnológico, hosting), completando "
     "las secciones vacías, eliminando duplicaciones, y agregando nuevas secciones (requerimientos, "
     "problemáticas, conclusiones)."),
]

for titulo, problema, solucion in problemas_tec:
    doc.add_heading(titulo, level=3)
    p = doc.add_paragraph()
    r = p.add_run("Problema: ")
    r.bold = True
    p.add_run(problema)
    p = doc.add_paragraph()
    r = p.add_run("Solución: ")
    r.bold = True
    p.add_run(solucion)


# ════════════════════════════════════════════
# 16. POSIBLES MEJORAS FUTURAS
# ════════════════════════════════════════════
doc.add_heading("16. Posibles Mejoras Futuras (Roadmap V2.0)", level=1)
doc.add_paragraph(
    "Con el objetivo de garantizar la escalabilidad del sistema, se proponen las "
    "siguientes expansiones:"
)
for m in [
    "Desarrollo de Aplicación Móvil nativa (Android/iOS).",
    "Integración de Inteligencia Artificial para resumir los apuntes de forma automática.",
    "Chat interno en tiempo real para consultas académicas entre estudiantes.",
    "Motor de recomendación automática de material complementario basado en materias cursadas.",
    "Sincronización e integración con calendarios escolares.",
    "Generación automática de cuestionarios (Quizzes) a partir del texto extraído de los apuntes.",
    "Integración de un servicio de email real (SMTP) para recuperación de contraseñas.",
    "Almacenamiento de archivos en la nube (AWS S3 o similar) para mayor durabilidad.",
    "Sistema de notificaciones push para avisar de nuevos apuntes en materias favoritas.",
]:
    doc.add_paragraph(m, style="List Bullet")


# ════════════════════════════════════════════
# 17. CONCLUSIONES
# ════════════════════════════════════════════
doc.add_heading("17. Conclusiones", level=1)
doc.add_paragraph(
    "El proyecto KIROKU demostró ser viable tanto desde el punto de vista técnico como "
    "organizativo. Se logró construir una plataforma web funcional que cumple con la "
    "mayoría de los requerimientos planteados en el contrato inicial, permitiendo a los "
    "estudiantes de la Escuela Técnica Fragata Libertad N°21 subir, organizar, buscar y "
    "compartir apuntes de clase de manera colaborativa."
)
doc.add_paragraph(
    "Desde el punto de vista técnico, se utilizó un stack moderno y liviano (Python/Flask + "
    "MySQL + JavaScript vanilla) que demostró ser adecuado para el alcance del proyecto. "
    "La arquitectura JWT permite una experiencia de usuario fluida sin recargar la página, "
    "y el sistema de roles y permisos garantiza la seguridad del contenido."
)
doc.add_paragraph(
    "Desde el punto de vista organizativo, el trabajo en equipo de 9 integrantes bajo "
    "metodología Scrum permitió dividir el proyecto en sprints manejables y entregar "
    "valor incrementally. Las dificultades de coordinación se resolvieron con herramientas "
    "de gestión (Trello) y comunicación constante (WhatsApp)."
)
doc.add_paragraph(
    "Las principales lecciones aprendidas fueron: la importancia de verificar la seguridad "
     "en cada funcionalidad (como el caso del bloqueo de usuarios), la necesidad de mantener "
     "la documentación actualizada con el código real, y el valor de la revisión de código "
     "para detectar vulnerabilidades XSS antes de que lleguen a producción."
)
doc.add_paragraph(
    "El proyecto queda en condiciones de ser desplegado en producción (Render + TiDB Cloud) "
    "y utilizado por los estudiantes de la escuela, con un camino claro para futuras mejoras "
    "como aplicación móvil, IA y notificaciones push."
)


# ════════════════════════════════════════════
# GUARDAR
# ════════════════════════════════════════════
output_dir = os.path.join(os.path.dirname(__file__), "Documentación")
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "Documentacion PI2.docx")
doc.save(output_path)
print(f"Documento generado: {output_path}")
