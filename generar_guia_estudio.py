from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, glob as g

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x26, 0x32, 0x38)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    hs.font.name = 'Calibri'

BLUE = RGBColor(0x1A, 0x23, 0x7E)
DARK = RGBColor(0x26, 0x32, 0x38)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GRAY = RGBColor(0x60, 0x60, 0x60)
RED = RGBColor(0xC6, 0x28, 0x28)

def add_qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run(f"P: {q}")
    r.bold = True
    r.font.color.rgb = BLUE
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"R: {a}")
    r2.font.color.rgb = DARK

def add_gloss(term, defn):
    p = doc.add_paragraph()
    r = p.add_run(f"{term}: ")
    r.bold = True
    r.font.color.rgb = BLUE
    r2 = p.add_run(defn)
    r2.font.color.rgb = DARK

# PORTADA
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("KIROKU")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Guia de Estudio para Defensa")
r.font.size = Pt(18)
r.font.color.rgb = DARK

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run("Escuela Tecnica Fragata Libertad N21\nProyecto Informatico 2 | Julio 2026")
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_page_break()

# INDICE
doc.add_heading("INDICE", level=1)
for item in [
    "1. Resumen del Proyecto",
    "2. Que decir en cada slide",
    "3. Preguntas del Profesor con Respuestas",
    "4. Glosario de Terminos Tecnicos",
    "5. Checklist Pre-Exposicion"
]:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. RESUMEN
doc.add_heading("1. RESUMEN DEL PROYECTO", level=1)

doc.add_heading("Que es KIROKU?", level=2)
doc.add_paragraph("KIROKU es una plataforma web colaborativa de apuntes para la Escuela Tecnica Fragata Libertad N21. Permite a los estudiantes subir, organizar, buscar y compartir apuntes de clase. Cuando un estudiante falta, consulta la plataforma en vez de pedir fotos por WhatsApp.")

p = doc.add_paragraph()
r = p.add_run('"Ningun estudiante se queda atras."')
r.bold = True
r.italic = True
r.font.color.rgb = BLUE
r.font.size = Pt(13)

doc.add_heading("Problema central", level=2)
doc.add_paragraph("No existe un repositorio unico, organizado y accesible donde los estudiantes de la Escuela Tecnica Fragata Libertad N21 puedan consultar los apuntes de clase cuando faltan. La recuperacion del material depende de contactar companeros de manera individual por WhatsApp.")

doc.add_heading("Problemas secundarios", level=2)
for item in [
    "Perdida o fuga de informacion por grupos de WhatsApp",
    "Apuntes incompletos, fragmentados o de mala calidad",
    "Desorganizacion general del material de estudio",
    "Desigualdad en el acceso al contenido academico",
    "Menor rendimiento academico de los estudiantes afectados"
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("KIROKU en numeros", level=2)
stats = [
    ("10 modulos implementados", ""),
    ("14 historias de usuario", "todas cumplidas"),
    ("17 archivos JavaScript", ""),
    ("12 modulos Python", ""),
    ("22 eventos de auditoria", ""),
    ("3 roles de usuario", "alumno, moderador, administrador"),
    ("Diseno responsive movil", ""),
    ("9 integrantes", ""),
    ("5 sprints de 2-4 dias cada uno", ""),
]
t = doc.add_table(rows=len(stats)+1, cols=2, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = "Metrica"
t.rows[0].cells[1].text = "Detalle"
for cell in t.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (m, d) in enumerate(stats):
    t.rows[i+1].cells[0].text = m
    t.rows[i+1].cells[1].text = d

doc.add_heading("Stack tecnologico", level=2)
stack = [
    ("Backend", "Python 3.14 + Flask", "API REST, rutas, autenticacion"),
    ("Base de datos", "MySQL / MariaDB", "XAMPP local + TiDB Cloud Serverless"),
    ("Autenticacion", "bcrypt + JWT", "Hash de contrasenas + tokens HTTP-only"),
    ("Frontend", "HTML5, CSS3, JavaScript vanilla", "Sin frameworks, diseno propio"),
    ("Graficos", "Chart.js v4", "Dashboard de estadisticas"),
    ("Iconos", "Lucide Icons (SVG)", "Renderizado consistente multiplataforma"),
    ("Audio", "Web Audio API", "Efectos de sonido generados en codigo"),
    ("Hosting", "Render", "Deploy con Gunicorn en produccion"),
    ("Control de versiones", "Git / GitHub", "Tags v1.0.0 a v1.0.10"),
]
t = doc.add_table(rows=len(stack)+1, cols=3, style='Table Grid')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Componente", "Tecnologia", "Detalle"]):
    t.rows[0].cells[j].text = h
    for p in t.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True
for i, (c, tech, det) in enumerate(stack):
    t.rows[i+1].cells[0].text = c
    t.rows[i+1].cells[1].text = tech
    t.rows[i+1].cells[2].text = det

doc.add_heading("Arquitectura", level=2)
doc.add_paragraph("Arquitectura Cliente-Servidor (Frontend/Backend) de tres capas:")
for capa in [
    "FRONTEND: HTML5 + CSS3 + JavaScript vanilla",
    "BACKEND: Python + Flask (12 modulos: auth, usuarios, cursos, materias, apuntes, valoraciones, busqueda, estadisticas, auditoria)",
    "BASE DE DATOS: MySQL / MariaDB"
]:
    doc.add_paragraph(capa, style='List Bullet')

doc.add_heading("Metodologia", level=2)
doc.add_paragraph("Scrum: 5 sprints de 2-4 dias cada uno")
doc.add_paragraph("Herramientas: Trello (tareas), WhatsApp (comunicacion), GitHub (codigo)")

doc.add_page_break()

# 2. QUE DECIR EN CADA SLIDE
doc.add_heading("2. QUE DECIR EN CADA SLIDE", level=1)

slides = [
    ("Slide 1 - Portada", "", "", "No se habla, se muestra."),
    ("Slide 2 - Agenda", "LEON", "", "Hoy vamos a presentar KIROKU en 6 partes: primero la problematica, luego la solucion KIROKU, despues la organizacion (equipo, Scrum y Trello), las historias de usuario y la demo, los aspectos tecnicos, y finalmente la reflexion y conclusiones."),
    ("Slide 3 - Problematica", "LEON", "", "Cuando un estudiante falta a clase, pierde el acceso a los contenidos trabajados. El problema central es que no existe un repositorio unico, organizado y accesible. La recuperacion del material depende de contactar companeros por WhatsApp. Esto genera perdida de informacion, apuntes incompletos, desorganizacion, desigualdad en el acceso, y menor rendimiento academico."),
    ("Slide 4 - Analisis de Usuarios", "MATIAS", "", "Tres tipos de usuarios: El Alumno que necesita consultar apuntes, subir material, y valorar. El Moderador que aprueba/rechaza contenido y gestiona materias. Y el Administrador que gestiona usuarios, roles, y ve estadisticas."),
    ("Slide 5 - Solucion: KIROKU", "LEON", "", "KIROKU es una plataforma web colaborativa para subir, organizar y compartir apuntes. Subida con drag and drop, clasificacion por materia/curso/fecha/autor, busqueda avanzada, valoracion con estrellas y Me Gusta, aprobacion por moderadores, panel de administracion, y auditoria. En numeros: 10 modulos, 14 HU, 17 JS, 12 Python, 22 auditorias, 3 roles, responsive movil."),
    ("Slide 6 - Objetivos", "MATIAS", "", "Objetivo general: disenar e implementar una plataforma que facilite el acceso democratico y organizado a los materiales de estudio. Especificos: centralizar apuntes, facilitar recuperacion ante inasistencias, fomentar companerismo, reducir dependencia de WhatsApp, mejorar organizacion academica."),
    ("Slide 7 - Organizacion del Equipo", "SANTINO", "", "9 integrantes con distintos niveles de dedicacion. Leon (Lider/Full Stack, full-time), Matias (SM/Sublider/Full Stack, full-time), Vito y Thiago (Frontend, partial-time), Ramiro (QA, partial-time), Santino (Documentador, full-time), Federico (Full Stack, full-time), Juan Pablo (QA, partial-time), Lucas (DBA/Full Stack, full-time)."),
    ("Slide 8 - Metodologia Scrum y Cronograma", "SANTINO", "", "5 sprints de 2-4 dias cada uno. Sprint 1: autenticacion y base. Sprint 2: gestion academica. Sprint 3: core del negocio (subida de archivos, aprobacion). Sprint 4: interaccion y busqueda. Sprint 5: admin y cierre (panel admin, Chart.js, auditoria, deploy)."),
    ("Slide 9 - Tablero Trello", "SANTINO", "", "Trello con 4 listas: Pendiente, En proceso, En revision, Finalizado. Cada integrante actualiza sus tarjetas. [Capturas del Trello]"),
    ("Slide 10 - Stack Tecnologico", "FEDERICO", "", "9 componentes: Python+Flask (backend), MySQL/MariaDB (BD), bcrypt+JWT (auth), HTML5+CSS3+JS vanilla (frontend), Chart.js (graficos), Lucide Icons SVG (iconos), Web Audio API (sonidos), Render+Gunicorn (hosting), Git+GitHub (versiones)."),
    ("Slide 11 - Arquitectura del Sistema", "LUCAS", "", "Cliente-Servidor de tres capas: Frontend (HTML5/CSS3/JS), Backend (Python+Flask, 12 modulos), Base de datos (MySQL/MariaDB). Comunicacion via HTTP y JSON."),
    ("Slide 12 - Seguridad y Auditoria", "JUAN PABLO", "", "6 capas: bcrypt+JWT HTTP-only, headers CSP/X-Frame-Options/X-Content-Type-Options, validacion de contrasena, escapeHtml() anti-XSS, bloqueo de usuarios, auditoria con 22 eventos."),
    ("Slide 13 - Historias de Usuario", "RAMIRO", "", "14 HU todas cumplidas: HU1-HU14 desde registro hasta personalizar perfil."),
    ("Slide 14 - Demo (titulo)", "", "", "Slide de transicion. No se habla."),
    ("Slide 15 - Demo: Modulo de Usuarios", "VITO", "", "Login con usuario o email, registro con validacion, recuperacion simulada, perfil con avatar y curso. [Capturas: app_login, app_home, app_perfil]"),
    ("Slide 16 - Demo: Modulo de Apuntes y Busqueda", "THIAGO", "", "Drag and drop, vista previa PDF/imagenes/video, busqueda avanzada, valoracion 1-5, Me Gusta, guardados. [Capturas: app_curso, app_apunte, app_busqueda]"),
    ("Slide 17 - Demo: Panel de Administracion y Estadisticas", "FEDERICO", "", "Gestion de usuarios (bloquear/activar/eliminar), graficos Chart.js, ranking colaboradores, responsive movil. [Capturas: app_admin, app_estadisticas, app_mobile]"),
    ("Slide 18 - Decisiones Tomadas", "RAMIRO", "", "6 decisiones: Python+Flask (experiencia del equipo), JS vanilla (control total), JWT HTTP-only (escalabilidad stateless), diseno post-it (identidad visual), remocion modo oscuro (incompleto/inconsistente), migracion emojis a Lucide SVG (renderizado consistente)."),
    ("Slide 19 - Reflexion Grupal", "TODOS", "1 frase c/u", "Leon: 'Liderar 9 personas me enseno que comunicar es mas importante que programar.' Matias: 'Scrum nos permitio convertir un proyecto enorme en sprints manejables.' Vito: 'Aprendi que el diseno no es solo como se ve, sino como se siente el usuario.' Thiago: 'El responsive design me enseno que no todos ven la misma pantalla que yo.' Ramiro: 'Tester no es buscar errores, es garantizar una buena experiencia.' Santino: 'Documentar me hizo entender que sin registro, el conocimiento se pierde.' Federico: 'Full stack es entender frontend, backend y como se conectan.' Juan Pablo: 'Las pruebas nos ahorraron problemas graves en produccion.' Lucas: 'La base de datos es el corazon del sistema: si falla, todo falla.'"),
    ("Slide 20 - Cierre", "", "", "No se habla. Portada de cierre igual a la Slide 1.")
]

for slide_name, who, role_note, text in slides:
    doc.add_heading(slide_name, level=2)
    if who:
        p = doc.add_paragraph()
        r = p.add_run(f"Quien habla: {who}")
        r.bold = True
        r.font.color.rgb = GREEN
    if role_note:
        p = doc.add_paragraph()
        r = p.add_run(f"Nota: {role_note}")
        r.bold = True
        r.font.color.rgb = RED
    if text:
        doc.add_paragraph(text)

doc.add_page_break()

# 3. PREGUNTAS
doc.add_heading("3. PREGUNTAS DEL PROFESOR CON RESPUESTAS", level=1)

doc.add_heading("Sobre la Problematica", level=2)
add_qa("Por que eligieron esta problematica y no otra?", "Es un problema que todos vivimos en la escuela. Cuando falta un companero, siempre pide fotos por WhatsApp y muchas veces no se las pasan o llegan incompletas.")
add_qa("Cual es el impacto real?", "Estudiantes que faltan quedan atras, reciben peores calificaciones, y se genera desigualdad en el acceso al contenido.")
add_qa("Ya existen soluciones similares?", "Google Classroom, Moodle, etc. Son genericas y complejas. KIROKU es simple, esta pensada para nuestro contexto, y tiene diseno estilo post-it.")
add_qa("Como saben que la van a usar?", "El problema es real. El diseno es simple. La barrera de entrada es baja: registrarse y unirse a un curso con codigo.")
add_qa("Que pasa si no quieren colaborar?", "No es obligatorio. Pero el sistema de valoraciones y ranking motiva a contribuir.")

doc.add_heading("Sobre la Solucion", level=2)
add_qa("Por que KIROKU?", "'Kiroku' significa 'registro' en japones.")
add_qa("Contenido inapropiado?", "Aprobacion obligatoria por moderadores. Nada se publica sin revision.")
add_qa("Cuantos usuarios puede manejar?", "Depende del hosting. JWT es stateless, permite escalar horizontalmente.")
add_qa("Si se cae el servidor?", "Render reinicia automaticamente. Datos seguros en TiDB Cloud con backups.")
add_qa("Que es lo mas innovador?", "Aprobacion por moderadores, diseno post-it, efectos de sonido con Web Audio API.")
add_qa("Mejoras futuras?", "Emails reales, app movil, IA para buscar por contenido, notificaciones push.")

doc.add_heading("Sobre Tecnologia", level=2)
add_qa("Por que Python y no Node.js?", "El equipo tenia mas experiencia. Flask es ligero y rapido.")
add_qa("Que es JWT?", "Token stateless: el servidor genera un token con la info del usuario. No guarda sesiones.")
add_qa("Que es bcrypt?", "Algoritmo de hash irrevertible para contrasenas.")
add_qa("HTTP-only que es?", "Cookie inaccesible desde JavaScript. Protege contra XSS.")
add_qa("Por que no React/Vue?", "JS vanilla fue suficiente. Sin dependencias, control total.")
add_qa("Que es escapeHtml?", "Reemplaza caracteres especiales por entidades HTML. Previene XSS.")
add_qa("Que es CORS?", "Controla que dominios pueden hacer peticiones a nuestro servidor.")
add_qa("Que es Gunicorn?", "Servidor HTTP para Python en produccion.")
add_qa("Por que TiDB Cloud?", "BD en la nube 24/7, compatible MySQL, tier gratuito.")
add_qa("Que es API REST?", "Interfaz con HTTP (GET/POST/PUT/DELETE) para comunicar frontend y backend.")
add_qa("Que es Web Audio API?", "API del navegador para generar sonidos en codigo.")
add_qa("Que es stateless?", "Sin estado: el servidor no guarda sesiones entre peticiones.")
add_qa("Headers CSP/X-Frame-Options?", "CSP controla scripts. X-Frame-Options previene clickjacking.")
add_qa("Que es una BD relacional?", "Tablas con filas y columnas, relacionadas por claves primarias y foraneas.")

doc.add_heading("Sobre Metodologia", level=2)
add_qa("Como se organizaron?", "Scrum con sprints de 2-4 dias. Trello con 4 listas. Comunicacion diaria por WhatsApp.")
add_qa("Que es Scrum?", "Metodologia agil con sprints cortos y entregas iterativas.")
add_qa("Si un integrante no cumple?", "Contacto directo, redistribucion de tareas, escalamiento al profesor si persiste.")
add_qa("Decisiones tecnicas?", "Consenso. El lider tenia la ultima palabra.")
add_qa("Tiempo por semana?", "10-15 horas por integrante.")
add_qa("Aprendizajes?", "Scrum, Git/GitHub, proyecto completo de software, resolver problemas reales.")
add_qa("Dificultades?", "Tecnicas: bugs, merge conflicts, JWT stale. Organizativas: coordinar 9 personas.")
add_qa("Importancia de la planificacion?", "Sin ella, 9 personas serian un caos. Scrum y Trello nos organizaron.")
add_qa("Que cambiarian?", "Mejores requisitos, roles mas claros, testing temprano, deploy antes.")
add_qa("Por que removieron el modo oscuro?", "Quedo incompleto e inconsistente. Preferimos consistencia a medias.")
add_qa("Por que migraron a Lucide SVG?", "Emojis se renderizaban distinto. Lucide garantiza consistencia.")

doc.add_heading("Preguntas Trampa", level=2)
add_qa("Mayor debilidad?", "Recuperacion de contrasena simulada. Hosting gratuito con limitaciones.")
add_qa("Apuntes falsos?", "Aprobacion por moderadores. Reportes de contenido.")
add_qa("Contrasenas seguras?", "Hash bcrypt irrevertible. Sin texto plano.")
add_qa("Datos personales?", "Minimizacion: solo nombre, email, contrasena. Sin datos sensibles.")
add_qa("Por que no NoSQL?", "Datos relacionales. MySQL ideal para JOINs.")
add_qa("Es escalable?", "Si. JWT stateless, BD escalable, codigo modular.")

doc.add_page_break()

# 4. GLOSARIO
doc.add_heading("4. GLOSARIO DE TERMINOS TECNICOS", level=1)

for term, defn in [
    ("API", "Interfaz de comunicacion entre sistemas. Usamos API REST."),
    ("Aprobacion de contenido", "Proceso donde un moderador revisa y decide publicar o rechazar un apunte."),
    ("Arquitectura Cliente-Servidor", "Modelo donde el navegador hace peticiones al servidor."),
    ("Auditoria", "Registro de eventos del sistema (login, subidas, etc.)."),
    ("Backend", "Parte del sistema que corre en el servidor."),
    ("bcrypt", "Algoritmo de hash de contrasenas irrevertible."),
    ("CORS", "Controla que dominios pueden acceder a nuestro servidor."),
    ("CRUD", "Create, Read, Update, Delete. Operaciones basicas sobre datos."),
    ("CSRF", "Ataque que envia peticiones falsas desde el navegador del usuario."),
    ("CSS", "Lenguaje para la apariencia visual de paginas web."),
    ("CSP", "Content Security Policy. Controla que scripts se ejecutan."),
    ("DBMS", "Sistema de gestion de bases de datos (MySQL/MariaDB)."),
    ("Deploy", "Publicar la aplicacion en un servidor de produccion."),
    ("DOM", "Representacion en memoria del HTML de una pagina."),
    ("Drag & Drop", "Arrastrar y soltar archivos con el mouse."),
    ("Endpoint", "Punto de acceso de una API. Ej: /api/login."),
    ("escapeHtml", "Reemplaza caracteres especiales por entidades HTML. Previene XSS."),
    ("ER", "Entidad-Relacion. Modelo para disenar bases de datos."),
    ("FK", "Foreign Key. Columna que referencia a otra tabla."),
    ("Flask", "Microframework web para Python."),
    ("Frontend", "Parte visible del sistema (HTML, CSS, JS)."),
    ("GET", "Metodo HTTP para obtener datos."),
    ("Gunicorn", "Servidor HTTP para Python en produccion."),
    ("Hash", "Resultado de una funcion matematica unidireccional."),
    ("Headers de seguridad", "Cabeceras HTTP: CSP, X-Frame-Options, X-Content-Type-Options."),
    ("HTML", "Lenguaje de marcado para estructura web."),
    ("HTTP", "Protocolo de comunicacion navegador-servidor."),
    ("HTTP-only", "Cookie inaccesible desde JavaScript."),
    ("JavaScript", "Lenguaje de programacion del navegador."),
    ("JSON", "Formato de intercambio de datos."),
    ("JWT", "JSON Web Token. Token de autenticacion stateless."),
    ("KPI", "Indicador clave de rendimiento."),
    ("Localhost", "Direccion del propio computador (127.0.0.1)."),
    ("Lucide Icons", "Biblioteca de iconos SVG consistentes."),
    ("MariaDB", "Fork de MySQL compatible."),
    ("Me Gusta", "Valoracion rapida sin escala numerica."),
    ("Metodo HTTP", "Accion sobre un recurso: GET, POST, PUT, DELETE."),
    ("Mockup", "Prototipo visual de la interfaz."),
    ("Modularizacion", "Organizar codigo en archivos separados por funcion."),
    ("MySQL", "Sistema de BD relacional con SQL."),
    ("Payload", "Datos dentro del JWT (ID, nombre, rol, expiracion)."),
    ("POST", "Metodo HTTP para crear recursos."),
    ("PyMySQL", "Libreria Python para MySQL."),
    ("Python", "Lenguaje de programacion del backend."),
    ("QA", "Quality Assurance. Aseguramiento de calidad."),
    ("Repositorio", "Lugar donde se guarda el codigo fuente (GitHub)."),
    ("Responsive", "Diseno adaptable a diferentes pantallas."),
    ("REST", "Estilo de arquitectura con HTTP sobre URLs."),
    ("RF", "Requerimiento Funcional."),
    ("RNF", "Requerimiento No Funcional."),
    ("Scrum", "Metodologia agil con sprints cortos."),
    ("Servidor", "Computador que ejecuta la aplicacion."),
    ("SQL", "Lenguaje para bases de datos relacionales."),
    ("Stateless", "Sin estado: sin sesiones guardadas en el servidor."),
    ("Trello", "Herramienta de gestion de tareas tipo Kanban."),
    ("Upload", "Subida de archivos al servidor."),
    ("URL", "Direccion de un recurso en internet."),
    ("Validator", "Funcion que verifica datos antes de guardarlos."),
    ("Vulnerabilidad", "Debilidad de seguridad explotable."),
    ("Web Audio API", "API del navegador para generar sonidos."),
    ("Wireframe", "Boceto visual sin estilo ni colores."),
    ("XAMPP", "Paquete Apache+MySQL+PHP para desarrollo local."),
    ("XSS", "Cross-Site Scripting. Inyeccion de JS malicioso."),
    ("X-Frame-Options", "Header que previene clickjacking."),
    ("X-Content-Type-Options", "Header que previene interpretacion MIME incorrecta.")
]:
    add_gloss(term, defn)

doc.add_page_break()

# 5. CHECKLIST
doc.add_heading("5. CHECKLIST PRE-EXPOSICION", level=1)
for c in [
    "Cada integrante sabe que slide le toca hablar",
    "Cada integrante estudio su seccion y puede responder preguntas",
    "Las capturas del Trello estan en la presentacion (slide 9)",
    "Las capturas de la app estan en la presentacion (slides 15-17)",
    "El PPTX se subio a Google Slides y se ve bien",
    "Hay una copia de backup en USB",
    "Navegador abierto con la plataforma funcionando",
    "Todos saben que es JWT, bcrypt, CRUD, Scrum, REST, escapeHtml",
    "Todos practicaron las frases de la reflexion",
    "Cada integrante sabe explicar las 6 decisiones (slide 18)",
    "Cada integrante conoce las 14 historias de usuario",
    "Cada integrante puede explicar los 10 modulos"
]:
    doc.add_paragraph(c, style='List Bullet')

# Guardar
matches = g.glob(r"C:\Users\rital\OneDrive\Desktop\COLEGIO\Proyecto Mitingay\Documen*\Guia_Estudio_KIROKU.docx")
if matches:
    output = matches[0]
else:
    import os
    doc_dir = [d for d in os.listdir(r"C:\Users\rital\OneDrive\Desktop\COLEGIO\Proyecto Mitingay") if d.startswith("Documen")][0]
    output = os.path.join(r"C:\Users\rital\OneDrive\Desktop\COLEGIO\Proyecto Mitingay", doc_dir, "Guia_Estudio_KIROKU.docx")
doc.save(output)
print(f"Guardado: {output}")
